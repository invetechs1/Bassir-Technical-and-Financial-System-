"""فحص شامل لنظام عزوم — يختبر كل نقطة نهاية وكل وظيفة (113 فحصاً).

التشغيل على الخادم بعد أي نشر أو تحديث:
    .venv/bin/python scripts/system_check.py
يعمل على قاعدة البيانات الحالية دون إتلافها، ويصلح أن يُشغَّل مرات متكررة.
"""
import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

RESULTS = []
def check(name, cond, detail=""):
    RESULTS.append((name, bool(cond), detail))
    mark = "✅" if cond else "❌"
    print(f"{mark} {name}" + (f"  — {detail}" if detail and not cond else ""))

# ---------- 1. تشغيل أول مرة: التهيئة والبذور ----------
from fastapi.testclient import TestClient
from app.main import app
from app.database import init_db
from app.auth import init_auth
from app.seed import seed_if_empty

init_db(); init_auth(); seed_if_empty()
c = TestClient(app)

# ---------- 2. المصادقة ----------
r = c.get("/api/status")
check("حماية API قبل الدخول (401)", r.status_code == 401)
r = c.get("/", follow_redirects=False)
check("تحويل الصفحة الرئيسية لتسجيل الدخول", r.status_code in (302, 307) and "/login" in r.headers.get("location", ""))
r = c.get("/login")
check("صفحة الدخول تعمل وبالهوية الخضراء", r.status_code == 200 and "AZOOM United Co." in r.text and "#175934" in r.text)
r = c.post("/api/login", json={"username": "azoom", "password": "wrong"})
check("رفض كلمة مرور خاطئة", r.status_code == 401)
r = c.post("/api/login", json={"username": "azoom", "password": "Azoom@2026"})
check("تسجيل الدخول", r.status_code == 200 and r.json()["user"]["username"] == "azoom")
r = c.get("/api/me")
check("/api/me", r.status_code == 200 and r.json()["username"] == "azoom")
r = c.post("/api/password", json={"old": "Azoom@2026", "new": "Test@12345"})
check("تغيير كلمة المرور", r.status_code == 200)
r = c.post("/api/login", json={"username": "azoom", "password": "Test@12345"})
check("الدخول بالكلمة الجديدة", r.status_code == 200)
r = c.post("/api/password", json={"old": "Test@12345", "new": "Azoom@2026"})
check("إرجاع كلمة المرور الأصلية", r.status_code == 200)
r = c.post("/api/password", json={"old": "Azoom@2026", "new": "short"})
check("رفض كلمة مرور قصيرة", r.status_code != 200)

# ---------- 3. الحالة والبذور ----------
r = c.get("/api/status"); st = r.json()
check("/api/status", r.status_code == 200 and st.get("ok"))
check(f"بذور الأسعار ({st.get('price_items')} بنداً)", st.get("price_items", 0) >= 600, str(st))
r = c.get("/api/proposals"); props = r.json()
refs = [p for p in props if p.get("is_reference") or "مرجعي" in str(p)]
check(f"العروض المرجعية المبذورة ({len(props)} عرضاً بالأرشيف)", len(props) >= 7)
r = c.get("/")
check("الواجهة الرئيسية بالهوية الجديدة (الشعار الرسمي + المجموعات)",
      "AZOOM" in r.text and "azoom-mark.png" in r.text and "nav-group" in r.text)
r = c.get("/static/app.js")
check("ملف app.js يُقدَّم", r.status_code == 200 and "makeReference" in r.text)
r = c.get("/static/styles.css")
check("ملف styles.css أخضر", r.status_code == 200 and "#175934" in r.text)

# ---------- 4. الإعدادات ----------
r = c.get("/api/settings"); s0 = r.json()
check("قراءة الإعدادات (بيانات عزوم الرسمية)",
      s0.get("company_cr") == "1010467099" and s0.get("company_vat_no") == "311917527400003"
      and s0.get("company_iban", "").startswith("SA49"))
r = c.put("/api/settings", json={"profit_pct": "18"})
r2 = c.get("/api/settings")
check("تعديل الإعدادات", r.status_code == 200 and r2.json()["profit_pct"] == "18")
c.put("/api/settings", json={"profit_pct": "15"})

# ---------- 5. قاعدة الأسعار ----------
r = c.get("/api/prices")
all_prices = r.json()
check(f"قائمة الأسعار ({len(all_prices)})", len(all_prices) >= 600)
r = c.get("/api/prices", params={"search": "دهان"})
check("البحث في الأسعار", r.status_code == 200 and len(r.json()) > 0)
r = c.post("/api/prices", json={"code": "TST-001", "category": "اختبار", "name": "بند اختبار الفحص", "unit": "م2", "unit_price": 100})
check("إضافة بند سعر", r.status_code == 200)
tid = r.json()["id"]
r = c.post("/api/prices", json={"id": tid, "code": "TST-001", "category": "اختبار", "name": "بند اختبار الفحص", "unit": "م2", "unit_price": 120})
r = c.get(f"/api/prices/{tid}/history")
check("تاريخ تغير السعر يُسجل", r.status_code == 200 and len(r.json()) >= 1)
r = c.get("/api/prices/export/csv")
check("تصدير CSV", r.status_code == 200 and "TST-001" in r.text)
csv_data = "code,category,name,unit,unit_price\nTST-002,اختبار,بند مستورد من CSV,عدد,55\n"
r = c.post("/api/prices/import/csv", files={"file": ("prices.csv", csv_data.encode("utf-8-sig"), "text/csv")})
check("استيراد CSV", r.status_code == 200)
r = c.get("/api/prices", params={"search": "TST-002"})
imported = r.json()
check("البند المستورد موجود", len(imported) == 1 and imported[0]["unit_price"] == 55)
r = c.delete(f"/api/prices/{tid}")
check("حذف بند سعر", r.status_code == 200)

# ---------- 6. المكتبة الفنية ----------
r = c.get("/api/library")
check(f"المكتبة الفنية ({len(r.json())} محتوى)", len(r.json()) >= 5)
track = [e for e in r.json() if "جازان" in e.get("body", "")]
check("سجل الخبرات الحقيقي (جازان 89 مليون...)", len(track) >= 1)
r = c.post("/api/library", json={"category": "اختبار", "title": "محتوى اختبار", "body": "نص", "tags": ""})
lid = r.json()["id"]
r = c.delete(f"/api/library/{lid}")
check("إضافة/حذف محتوى مكتبة", r.status_code == 200)

# ---------- 7. وثائق الشركة ----------
r = c.get("/api/docs"); docs = r.json()
check(f"وثائق الشركة ({len(docs)})", len(docs) >= 9)
check("حساب حالة الصلاحية", all("status" in d for d in docs))
official = [d for d in docs if d.get("number")]
check(f"الوثائق بأرقامها الرسمية ({len(official)})", len(official) >= 4)

# ---------- 8. توليد عرض متكامل (محرك القوالب + التشابه + المستودع) ----------
from openpyxl import Workbook
wb = Workbook(); ws = wb.active
ws.append(["م", "وصف الأعمال", "الوحدة", "الكمية", "سعر الوحدة", "الإجمالي"])
ws.append([1, "توريد وتركيب عزل مائي للأسطح", "م2", 500, 40, 20000])
ws.append([2, "أعمال دهانات خارجية عازلة للحرارة", "م2", 1000, 25, 25000])
buf = io.BytesIO(); wb.save(buf)

r = c.post("/api/proposals/generate",
           data={"title": "مشروع عزل وترميم مبانٍ حكومية بمنطقة حائل", "client": "أمانة منطقة حائل", "entity_type": "government"},
           files=[("files", ("boq.xlsx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"))])
check("توليد عرض جديد", r.status_code == 200, r.text[:200])
prop = r.json(); pid = prop["id"]; data = prop["data"]
check("أقسام فنية (11 قسماً)", len(data.get("technical_sections", [])) >= 10)
check("جدول كميات غير فارغ", len(data.get("boq", [])) > 0)
check("خطة تنفيذية بمراحل", len(data.get("plan", [])) >= 3)
check("مصفوفة الالتزام", len(data.get("compliance_matrix", [])) >= 3)
check("التشابه وجد عروضاً مرجعية (عزل حائل/المطارات)", len(data.get("similar_refs", [])) >= 1,
      str(data.get("similar_refs")))

fin = data.get("financial", {})
direct = round(sum(l["total"] for l in data["boq"]), 2)
ok_math = (abs(fin["direct_cost"] - direct) < 1
           and abs(fin["overhead"] - direct * 0.12) < 1
           and abs(fin["risk"] - direct * 0.03) < 1
           and abs(fin["subtotal"] - (direct + fin["overhead"] + fin["risk"] + fin["profit"])) < 1
           and abs(fin["vat"] - fin["subtotal"] * 0.15) < 1
           and abs(fin["grand_total"] - (fin["subtotal"] + fin["vat"])) < 1)
check("صحة الحسابات المالية (إداري 12% + مخاطر 3% + ربح + ضريبة 15%)", ok_math, str(fin))

r = c.get("/api/proposals/similar", params={"q": "مشروع عزل أسطح ومعالجة تسربات لمطارات"})
check("نقطة نهاية التشابه", r.status_code == 200 and len(r.json()) >= 1)

r = c.get(f"/api/proposals/{pid}")
check("قراءة عرض واحد", r.status_code == 200)
r = c.put(f"/api/proposals/{pid}", json={"status": "won"})
check("تحديث حالة العرض (فائز)", r.status_code == 200)

# ---------- 9. التصدير Word + Excel ----------
_titles = [sec["title"] for sec in data.get("technical_sections", [])]
check("بنية العرض الحقيقية (خطاب التقديم أولاً + معلومات الشركة + من نحن)",
      _titles[:1] == ["خطاب التقديم"] and "معلومات الشركة" in _titles and "من نحن" in _titles,
      str(_titles[:5]))
_bodies_all = " ".join(sec["body"] for sec in data.get("technical_sections", []))
check("لا ملاحظات تحريرية في نصوص العميل", "حرّر هذا القسم" not in _bodies_all)
r = c.get(f"/api/proposals/{pid}/export/docx")
check("تصدير Word", r.status_code == 200 and len(r.content) > 10000)
from docx import Document
d = Document(io.BytesIO(r.content))
footer_ok = any("+966114880122" in p.text and "1010467099" in p.text
                for s in d.sections for p in s.footer.paragraphs)
check("تذييل Word الرسمي في كل صفحة", footer_ok)
_doc_text = "\n".join(p2.text for p2 in d.paragraphs)
for _t2 in d.tables:
    for _row in _t2.rows:
        _doc_text += "\n" + " ".join(cell.text for cell in _row.cells)
check("Word للعميل: لا هامش ربح ولا مخاطر ولا مصاريف إدارية",
      all(x not in _doc_text for x in ("هامش الربح", "احتياطي المخاطر", "المصاريف الإدارية", "التكلفة المباشرة")))
from app.proposal_builder import client_facing_pricing as _cfp
_loaded = _cfp(data["boq"], fin)
check("الأسعار المحمَّلة: مجموع البنود = الإجمالي قبل الضريبة",
      abs(round(sum(l["total"] for l in _loaded), 2) - fin["subtotal"]) < 0.01)
r = c.get(f"/api/proposals/{pid}/export/xlsx")
check("تصدير Excel", r.status_code == 200 and len(r.content) > 3000)
from openpyxl import load_workbook
wb2 = load_workbook(io.BytesIO(r.content))
check("ملف Excel سليم", len(wb2.sheetnames) >= 1)
_xl_text = " ".join(str(c2.value) for _r2 in wb2.active.iter_rows() for c2 in _r2 if c2.value)
check("Excel للعميل: لا نسب داخلية",
      all(x not in _xl_text for x in ("هامش الربح", "احتياطي المخاطر", "المصاريف الإدارية")))

# ---------- 10. المستودع المعرفي ----------
boq_txt = """عرض مالي — مشروع إنشاء حدائق وممرات مشاة
أعمال زراعة أشجار وشجيرات محلية
عدد 300 150.00 45,000.00
توريد وتركيب مظلات خشبية للجلسات
عدد 20 3,500.00 70,000.00
""".encode()
r = c.post("/api/repo/upload",
           data={"source_type": "عرض فني سابق", "company": "شركة الفحص", "notes": "فحص عميق", "as_reference": "1"},
           files=[("files", ("عرض حدائق وممرات.txt", boq_txt, "text/plain"))])
rec = r.json()[0]
check("رفع للمستودع + استخراج بندين", r.status_code == 200 and rec["items_count"] == 2, str(rec))
check("إنشاء عرض مرجعي تلقائي من الملف", bool(rec.get("reference")) or "مسبقاً" in str(rec.get("reference_note", "")), str(rec))
rfid = rec["id"]
r = c.post(f"/api/repo/{rfid}/make-reference")
check("منع تكرار المرجعي (409)", r.status_code == 409)
r = c.get("/api/repo")
check("قائمة المستودع + الإحصاءات", r.status_code == 200 and r.json()["stats"]["market_items"] >= 2)
r = c.get("/api/market/search", params={"q": "مظلات"})
mk = r.json()
check("مقارنة أسعار السوق", r.status_code == 200 and mk["benchmark"]["count"] >= 1 and mk["benchmark"]["avg"] == 3500)
r = c.delete(f"/api/repo/{rfid}")
check("حذف ملف من المستودع", r.status_code == 200)

# ملف مصور بلا نص — تشخيص واضح لا انهيار
r = c.post("/api/repo/upload", data={"source_type": "عرض شركة منافسة", "company": "", "notes": "", "as_reference": ""},
           files=[("files", ("scan.pdf", b"%PDF-1.4 fake", "application/pdf"))])
check("ملف مصور: تشخيص واضح بلا خطأ", r.status_code == 200 and r.json()[0].get("note"))

# ---------- 11. محلل الفرص ----------
risk_text = ("اشتراطات التأهيل: تصنيف مقاولين درجة ثانية في المباني، خبرة لا تقل عن 10 سنوات، "
             "ضمان نهائي 10% وغرامات تأخير 20%، نسبة محتوى محلي 40%").encode()
r = c.post("/api/opportunity", data={"title": "منافسة صيانة وتشغيل مباني بلدية", "client": "بلدية"},
           files=[("files", ("terms.txt", risk_text, "text/plain"))])
opp = r.json()
check("محلل الفرص: درجة ونتيجة", r.status_code == 200 and 0 <= opp.get("score", -1) <= 100 and opp.get("verdict"))
check("رصد مخاطر التأهيل (تصنيف/غرامات/محتوى محلي)", len(opp.get("qualification_warnings", [])) >= 2,
      str(opp.get("qualification_warnings")))
check("عوامل التقييم الخمسة", len(opp.get("factors", [])) == 5)

# ---------- 12. التحليلات ----------
r = c.get("/api/analytics"); an = r.json()
check("التحليلات: نسبة الفوز والقيم", r.status_code == 200 and "win_rate" in an.get("totals", {}) and an["totals"].get("won_value", 0) > 0)

# ---------- 13. اعتماد (الشبكة محجوبة هنا — يجب أن يفشل بلطف) ----------
r = c.post("/api/etimad/fetch", params={"pages": 1})
graceful = r.status_code in (200, 502, 503) and (r.status_code != 200 or "error" in str(r.json()) or isinstance(r.json(), dict))
check("جلب اعتماد يفشل بلطف عند حجب الشبكة", graceful, r.text[:150])
r = c.get("/api/etimad")
check("قائمة منافسات اعتماد", r.status_code == 200)

# ---------- 13ب. مشاريع منصة فرصة ----------
r = c.get("/api/forsah")
check("قائمة مشاريع فرصة + التصنيفات",
      r.status_code == 200 and len(r.json().get("categories", [])) == 5)
_saved_fs = c.get("/api/settings").json()
r = c.post("/api/forsah/fetch")
_j = r.json()
if not _saved_fs.get("forsah_email"):
    check("سحب فرصة بلا بيانات دخول: رسالة إرشادية", r.status_code == 200 and not _j.get("ok")
          and "بريد" in _j.get("error", ""), str(_j))
else:
    check("سحب فرصة: استجابة سليمة (نجاح أو تشخيص واضح)",
          r.status_code == 200 and (_j.get("ok") or _j.get("error")), str(_j)[:150])
r = c.get("/")
check("صفحة فرصة في الواجهة", "مشاريع منصة فرصة" in r.text and 'id="page-forsah"' in r.text)

# ---------- 14. حذف العرض + الخروج ----------
r = c.delete(f"/api/proposals/{pid}")
check("حذف عرض", r.status_code == 200)
r = c.post("/api/logout")
check("تسجيل الخروج", r.status_code == 200)
r = c.get("/api/status")
check("انقطاع الوصول بعد الخروج", r.status_code == 401)

# ---------- 15. تصنيف القطاعات (حكومي/خاص/صندوق الاستثمارات/مطارات) ----------
c.post("/api/login", json={"username": "azoom", "password": "Azoom@2026"})
apt_txt = """عرض مالي — أعمال صيانة مدرجات مطار
إعادة تأهيل طبقات الأسفلت لمدرج الطائرات
م2 4,000 85.00 340,000.00
دهانات علامات أرضية للمدرجات وممرات الطائرات
م.ط 6,000 12.50 75,000.00
""".encode()
r = c.post("/api/repo/upload",
           data={"source_type": "عرض عزوم سابق", "company": "", "notes": "", "as_reference": "1", "sector": "airports"},
           files=[("files", ("عرض صيانة مدرجات.txt", apt_txt, "text/plain"))])
rec_a = r.json()[0]
check("رفع بقطاع مطارات + مرجعي", r.status_code == 200 and rec_a["items_count"] == 2
      and (rec_a.get("reference") or "مسبقاً" in str(rec_a.get("reference_note", ""))), str(rec_a))
r = c.get("/api/market/search", params={"q": "مدرج", "sector": "airports"})
check("بحث السوق مصفّى بقطاع المطارات", r.json()["benchmark"]["count"] >= 1)
r = c.get("/api/market/search", params={"q": "مدرج", "sector": "private"})
check("التصفية تستبعد القطاعات الأخرى", r.json()["benchmark"]["count"] == 0)
r = c.get("/api/proposals/similar", params={"q": "مشروع أعمال عزل أسطح ومباني", "sector": "airports"})
sims = r.json()
check("التشابه يقدّم عروض نفس القطاع أولاً", len(sims) >= 1 and sims[0].get("sector_match") is True, str(sims[:2]))
r = c.get("/api/analytics")
check("التحليلات تشمل القطاعات الأربعة", set(r.json()["by_entity"]) == {"government", "private", "pif", "airports"})
r = c.get("/api/repo")
sect_file = [f for f in r.json()["files"] if f["filename"] == "عرض صيانة مدرجات.txt"]
check("القطاع محفوظ على ملف المستودع", sect_file and sect_file[0].get("sector") == "airports")

# ---------- 16. إعادة تشغيل: البذور لا تتكرر ----------
n_before = c.get("/api/status").json() if False else None
c.post("/api/login", json={"username": "azoom", "password": "Azoom@2026"})
before = c.get("/api/status").json()["price_items"]
seed_if_empty()
after = c.get("/api/status").json()["price_items"]
check("إعادة التشغيل لا تكرر البذور", before == after, f"{before} -> {after}")

# ---------- 15ب. المستودع الفني ومحرك الأسلوب ----------
c.post("/api/login", json={"username": "azoom", "password": "Azoom@2026"})
tech_text = """نبذة عن الشركة
تُعد شركة عزوم المتحدة للمقاولات من الشركات الرائدة في تنفيذ مشاريع المباني الحكومية، وقد نفذت الشركة مشاريع عزل وترميم وتشطيب تجاوزت قيمتها أربعمائة مليون ريال خلال السنوات الماضية بمعدلات إنجاز موثقة.
نطاق العمل
يشمل نطاق العمل تنفيذ أعمال العزل المائي والحراري للأسطح والمباني، وأعمال الترميم الإنشائي، ومعالجة التشققات، وفق كراسة الشروط والمواصفات الفنية المعتمدة من الجهة المالكة.
منهجية التنفيذ
تعتمد الشركة منهجية تنفيذ مرحلية تبدأ بالتعبئة واستلام الموقع، ثم تنفيذ الأعمال ببرنامج زمني معتمد من الاستشاري، مع رفع مستخلصات شهرية موثقة بنسب الإنجاز الفعلية في الموقع.
خطة ضمان الجودة
تطبق الشركة نظام إدارة جودة يشمل فحص المواد قبل التوريد، واختبارات معتمدة من مختبرات مستقلة، وسجلات تفتيش يومية يعتمدها الاستشاري قبل الانتقال لأي مرحلة لاحقة.
خطة السلامة والصحة المهنية
تلتزم الشركة بتطبيق اشتراطات السلامة في مواقع العمل، وتوفير مهمات الوقاية الشخصية لكافة العاملين، وتعيين مشرف سلامة متفرغ للموقع طوال مدة التنفيذ."""

for i in range(3):
    r = c.post("/api/repo/upload",
               data={"source_type": "عرض فني سابق", "company": "عزوم", "notes": "فحص أسلوب",
                     "as_reference": "", "sector": "government"},
               files=[("files", (f"عرض فني عزل مباني {i+1}.txt",
                                 tech_text.encode(), "text/plain"))])
    assert r.status_code == 200, r.text
tech_rec = r.json()[0]
check("رفع عرض فني → استخراج أقسام وفقرات", bool(tech_rec.get("tech", {}).get("sections", 0) >= 4),
      str(tech_rec.get("tech")))
r = c.post("/api/style-profile/rebuild")
sp = r.json()
check("استخلاص بصمة الكتابة", r.status_code == 200 and sp.get("sample_count", 0) >= 10, str(sp)[:150])
r = c.get("/api/paragraph-bank")
bank_items = r.json()
check(f"بنك الفقرات المعتمد ({len(bank_items)})", len(bank_items) >= 10)
r = c.get("/api/repository/technical")
check("قائمة المستودع الفني", r.status_code == 200 and len(r.json()["documents"]) >= 3)

r = c.post("/api/proposals/generate",
           data={"title": "مشروع عزل وترميم مبانٍ حكومية بالقصيم", "client": "أمانة القصيم",
                 "entity_type": "government"})
sty = r.json()["data"].get("style", {})
spid = r.json()["id"]
check("العرض مبني من بنك الفقرات (bank_ratio > 0)",
      sty.get("bank_ratio", 0) > 0 and sty.get("score", 0) > 0, str(sty))
_bodies = " ".join(sec["body"] for sec in r.json()["data"]["technical_sections"])
check("لا عبارات من القائمة السوداء في المخرَج",
      not any(b in _bodies for b in ("حلول مبتكرة", "شريك النجاح", "في الختام", "نسعى جاهدين")))
c.delete(f"/api/proposals/{spid}")
if bank_items:
    r = c.put(f"/api/paragraphs/{bank_items[0]['id']}", json={"approved": False})
    check("سحب اعتماد فقرة", r.status_code == 200)
    c.put(f"/api/paragraphs/{bank_items[0]['id']}", json={"approved": True})

# ---------- 15ج. تصنيف فئة المشروع (إنشاءات / صيانة وتشغيل / ...) ----------
from app.style_engine import detect_project_kind as _dpk
check("المصنف: صيانة وتشغيل", _dpk("عقد صيانة وتشغيل المرافق مع الصيانة الوقائية ومعالجة الأعطال")[0] == "صيانة وتشغيل")
check("المصنف: إنشاءات", _dpk("إنشاء وتشطيب مبنى خرسانة مسلحة مع أعمال العزل")[0] == "إنشاءات وتشطيبات")
check("المصنف: نظافة", _dpk("خدمات نظافة وتنظيف شاملة مع مكافحة الحشرات")[0] == "نظافة")

_maint_text = """نطاق العمل
يشمل نطاق العمل تشغيل وصيانة الأنظمة الكهربائية والميكانيكية، وبرامج الصيانة الوقائية الدورية، ومعالجة الأعطال والبلاغات ضمن أزمنة الاستجابة المحددة بالعقد المعتمد.
منهجية التنفيذ
نبدأ باستلام المواقع وجرد الأصول، ثم تنفيذ برنامج الصيانة الوقائية الشهري بسجلات لكل أصل، ومعالجة البلاغات بفريق مناوب، وتقارير شهرية بمؤشرات الأداء المعتمدة.""".encode()
r = c.post("/api/repo/upload",
           data={"source_type": "عرض فني سابق", "company": "عزوم", "notes": "فحص فئة",
                 "as_reference": "", "sector": "government"},
           files=[("files", ("عرض صيانة فحص الفئات.txt", _maint_text, "text/plain"))])
check("رفع عرض صيانة → تصنيف تلقائي في المستودع الفني", r.status_code == 200)
_docs = c.get("/api/repository/technical").json()["documents"]
_md = next((dd for dd in _docs if "صيانة فحص الفئات" in dd["filename"]), None)
check("فئة المستند المكتشفة = صيانة وتشغيل", _md and _md["project_kind"] == "صيانة وتشغيل", str(_md))

r = c.post("/api/proposals/generate",
           data={"title": "مشروع صيانة وتشغيل مرافق تعليمية", "client": "إدارة تعليم", "entity_type": "government"})
_dk = r.json()["data"]
check("توليد صيانة: الفئة مكتشفة والخطة تشغيلية سنوية",
      _dk.get("project_kind") == "صيانة وتشغيل" and "التشغيل والصيانة الدورية" in _dk["plan"][1]["phase"],
      str((_dk.get("project_kind"), _dk["plan"][1]["phase"])))
_bank_refs = {s3.get("source_ref", "") for s3 in _dk["technical_sections"] if s3.get("source") == "bank"}
check("بنك الفقرات اختار عروض الصيانة لا الإنشاءات",
      any("صيانة" in (ref or "") for ref in _bank_refs), str(_bank_refs))
c.delete(f"/api/proposals/{r.json()['id']}")

# ---------- 16. تعدد الشركات والأدوار ----------
c.post("/api/login", json={"username": "azoom", "password": "Azoom@2026"})
me = c.get("/api/me").json()
check("me: الدور والشركة ومدير المنصة", me.get("role") == "owner" and me.get("company_id") == 1
      and me.get("is_platform_admin") and me.get("is_admin"), str(me))
r = c.get("/api/me/companies")
check("قائمة شركات المستخدم", r.status_code == 200 and len(r.json()) >= 1)
azoom_prices_before = c.get("/api/status").json()["price_items"]

# شركة معزولة للاختبار (قابلة للتكرار: 409 عند إعادة التشغيل)
r = c.post("/api/companies", json={"name": "شركة الفحص المعزولة", "plan": "trial",
                                   "owner_username": "isocheck", "owner_password": "Iso@12345"})
check("إنشاء شركة جديدة بمالكها (أو موجودة من فحص سابق)", r.status_code in (200, 409), r.text[:120])

c2 = TestClient(app)
r = c2.post("/api/login", json={"username": "isocheck", "password": "Iso@12345"})
check("دخول مالك الشركة الثانية", r.status_code == 200)
me2 = c2.get("/api/me").json()
check("سياق الشركة الثانية (اسم وخطة ودور)",
      me2.get("company_name") == "شركة الفحص المعزولة" and me2.get("role") == "owner"
      and me2.get("company_id") != 1, str(me2))
st2 = c2.get("/api/status").json()
check("العزل: لا أسعار ولا عروض من عزوم في الشركة الثانية",
      st2["price_items"] == 0 or st2["price_items"] < 5, str(st2))
check("العزل: أرشيف عروض فارغ", len(c2.get("/api/proposals").json()) == 0)
s2 = c2.get("/api/settings").json()
check("العزل: إعدادات مستقلة (لا آيبان عزوم)",
      s2.get("company_name") == "شركة الفحص المعزولة" and not s2.get("company_iban"))
r = c2.post("/api/prices", json={"code": "LIT-LB.1", "category": "اختبار عزل",
                                 "name": "بند بكود مكرر عبر الشركات", "unit": "م2", "unit_price": 77})
check("نفس كود البند مسموح عبر شركتين", r.status_code == 200, r.text[:120])
check("وعدد أسعار الشركة الثانية = 1", len(c2.get("/api/prices").json()) == 1)
r = c2.post("/api/session/company/1")
check("تبديل لشركة بلا عضوية → 403", r.status_code == 403)
r = c2.get("/api/paragraph-bank")
check("العزل: بنك فقرات الشركة الثانية محجوب (402 للتجريبي) أو فارغ",
      r.status_code == 402 or (r.status_code == 200 and len(r.json()) == 0), r.text[:100])

# دور المشاهد في عزوم
r = c.post("/api/members", json={"username": "viewcheck", "password": "View@12345", "role": "viewer"})
check("دعوة مُشاهد لشركة عزوم", r.status_code == 200, r.text[:120])
c3 = TestClient(app)
r = c3.post("/api/login", json={"username": "viewcheck", "password": "View@12345"})
check("دخول المُشاهد", r.status_code == 200)
check("المُشاهد محجوب عن قاعدة الأسعار (403)", c3.get("/api/prices").status_code == 403)
check("المُشاهد محجوب عن التحليلات (403)", c3.get("/api/analytics").status_code == 403)
check("المُشاهد يقرأ لوحة التحكم", c3.get("/api/proposals").status_code == 200)
check("المُشاهد لا يكتب (403)",
      c3.put("/api/settings", json={"profit_pct": "99"}).status_code == 403)
check("المُشاهد لا يولد عروضاً (403)",
      c3.post("/api/proposals/generate", data={"title": "x", "client": "y"}).status_code == 403)
check("المُشاهد لا يدير الأعضاء (403)", c3.get("/api/members").status_code == 403)
check("بيانات عزوم سليمة بعد كل الفحوص",
      c.get("/api/status").json()["price_items"] == azoom_prices_before)

# ---------- 17. الاشتراكات والتسجيل الذاتي ----------
r = c.post("/api/signup", json={"name": "شركة التسجيل الذاتي للفحص", "cr_no": "9990001112",
                                "owner_username": "signupcheck", "owner_password": "Sign@12345",
                                "sector": "المقاولات"})
check("التسجيل الذاتي (شركة تجريبية 14 يوماً) أو 409 عند التكرار",
      r.status_code in (200, 409), r.text[:120])
r = c.post("/api/signup", json={"name": "شركة أخرى", "cr_no": "9990001112",
                                "owner_username": "someoneelse", "owner_password": "Else@12345"})
check("سجل تجاري مكرر → 409", r.status_code == 409)

c4 = TestClient(app)
r = c4.post("/api/login", json={"username": "signupcheck", "password": "Sign@12345"})
check("دخول مالك الشركة المسجلة ذاتياً", r.status_code == 200)
me4 = c4.get("/api/me").json()
check("خطة تجريبية للشركة الجديدة", me4.get("plan") == "trial", str(me4)[:120])

# بوابات الميزات: التجريبي بلا اعتماد/فرصة ولا محرك أسلوب → 402
check("بوابة الميزات: اعتماد 402 للخطة التجريبية", c4.get("/api/etimad").status_code == 402)
check("بوابة الميزات: بصمة الكتابة 402 للخطة التجريبية",
      c4.get("/api/style-profile").status_code == 402)
check("عزوم (مؤسسي) تصل لاعتماد طبيعياً", c.get("/api/etimad").status_code == 200)

# دورة الحياة: تجربة منتهية → قراءة وتصدير فقط (402 على الكتابة)
from app.database import get_db as _gdb
with _gdb() as _db:
    from datetime import datetime, timedelta, timezone
    _ended = (datetime.now(timezone.utc) - timedelta(days=5)).isoformat()
    _db.execute("UPDATE companies SET trial_ends_at = ? WHERE cr_no = '9990001112'", (_ended,))
check("تجربة منتهية: القراءة تعمل", c4.get("/api/proposals").status_code == 200)
r = c4.post("/api/proposals/generate", data={"title": "x", "client": "y"})
check("تجربة منتهية: الكتابة 402", r.status_code == 402 and "التجربة" in r.json()["detail"], r.text[:120])
with _gdb() as _db:
    _db.execute("UPDATE companies SET trial_ends_at = '2030-01-01T00:00:00+00:00' "
                "WHERE cr_no = '9990001112'")

# الفوترة: شركة مدفوعة تُفوتر مرة واحدة للفترة، والتجارب تُتخطى
with _gdb() as _db:
    _db.execute("UPDATE companies SET plan = 'basic' WHERE cr_no = '9990001112'")
r1 = c.post("/api/platform/invoices/issue").json()
r2 = c.post("/api/platform/invoices/issue").json()
check("إصدار الفواتير الشهرية يعمل وبلا تكرار للفترة",
      (r1["issued"] + r1["skipped"]) >= 1 and r2["issued"] == 0, f"{r1} ثم {r2}")
with _gdb() as _db:
    _db.execute("UPDATE companies SET plan = 'trial' WHERE cr_no = '9990001112'")
r = c.get("/api/platform/metrics")
check("مؤشرات مدير المنصة", r.status_code == 200 and "mrr" in r.json())
check("المؤشرات محجوبة عن غير مدير المنصة", c2.get("/api/platform/metrics").status_code == 403)
r = c.get("/signup")
check("صفحة التسجيل الذاتي العامة", r.status_code == 200 and "تجربة مجانية" in r.text)

# ---------- 18. تقسيم البند إلى أجزاء مستقلة (كمية ووحدة وسعر لكل جزء) ----------
r = c.post("/api/proposals/generate", data={"title": "فحص تقسيم بنود جدول الكميات",
                                            "client": "جهة الفحص", "entity_type": "government"})
check("توليد عرض لفحص تقسيم البنود", r.status_code == 200, r.text[:120])
_sp = r.json()["id"]
_sd = c.get(f"/api/proposals/{_sp}").json()["data"]
_sd["boq"][0]["children"] = [
    {"name": "توريد المواد", "unit": "طن", "qty": 5, "unit_price": 1000},
    {"name": "أعمال التركيب", "unit": "م2", "qty": 20, "unit_price": 150},
    {"name": "الاختبار والتشغيل", "unit": "مقطوعية", "qty": 1, "unit_price": 2500},
]
c.put(f"/api/proposals/{_sp}", json={"data": _sd})
_sd2 = c.get(f"/api/proposals/{_sp}").json()["data"]
_sl = _sd2["boq"][0]
check("كل جزء: الإجمالي = كميته × سعره",
      all(abs(k["total"] - k["qty"] * k["unit_price"]) < 0.01 for k in _sl["children"])
      and _sl["children"][0]["total"] == 5000 and _sl["children"][1]["total"] == 3000,
      str(_sl["children"])[:200])
check("إجمالي البند الأب = مجموع أجزائه (10500)", abs(_sl["total"] - 10500) < 0.01, str(_sl["total"]))
check("التكلفة المباشرة متسقة بعد التقسيم",
      abs(_sd2["financial"]["direct_cost"] - sum(l["total"] for l in _sd2["boq"])) < 0.05)
from app.proposal_builder import client_facing_pricing as _cfp
_cf = _cfp(_sd2["boq"], _sd2["financial"])
check("التسعير المحمَّل: مجموع الأجزاء المحمَّلة = إجمالي الأب",
      abs(_cf[0]["total"] - round(sum(k["total"] for k in _cf[0]["children"]), 2)) < 0.01)
check("التسعير المحمَّل: مجموع البنود = الإجمالي قبل الضريبة",
      abs(round(sum(l["total"] for l in _cf), 2) - _sd2["financial"]["subtotal"]) < 0.05)
r = c.get(f"/api/proposals/{_sp}/export/xlsx")
_rows = []
if r.status_code == 200:
    _wbx = load_workbook(io.BytesIO(r.content))
    for _wsx in _wbx.worksheets:
        for _row in _wsx.iter_rows(values_only=True):
            _rows.append(" | ".join(str(x) for x in _row if x is not None))
check("إكسل: صف الجزء يحمل كميته ووحدته (طن × 5)",
      any("توريد المواد" in x and "طن" in x and (" 5 " in f" {x} " or "| 5 |" in x) for x in _rows),
      str([x for x in _rows if "توريد" in x])[:200])
# توافق رجعي: جزء قديم بلا كمية/وحدة يرث كمية الأب ووحدته
_sd2["boq"][0]["children"] = [{"name": "جزء قديم", "unit_price": 100}]
c.put(f"/api/proposals/{_sp}", json={"data": _sd2})
_sl3 = c.get(f"/api/proposals/{_sp}").json()["data"]["boq"][0]
check("توافق رجعي: الجزء بلا كمية يرث كمية الأب ووحدته",
      _sl3["children"][0]["qty"] == _sl3["qty"] and _sl3["children"][0]["unit"] == _sl3["unit"])
c.delete(f"/api/proposals/{_sp}")

# ---------- 19. تنفيذ المشاريع: مقاولو الباطن وقفل تقارير الإنتاجية ----------
r = c.post("/api/execution/subcontractors", json={"name": "مؤسسة فحص الدرع", "trade": "خرسانة"})
check("تسجيل مقاول باطن (أو 409 لتكرار الاسم)", r.status_code in (200, 409), r.text[:120])
_opts = c.get("/api/execution/executors").json()
check("القائمة المنسدلة: عمالة الشركة أولاً ثم مقاولو الباطن",
      _opts and _opts[0]["label"] == "عمالة الشركة"
      and any("فحص الدرع" in o["label"] for o in _opts))
r = c.post("/api/execution/projects", json={"name": "مشروع فحص التنفيذ", "client": "جهة الفحص"})
_xp = r.json()["id"]
r = c.post("/api/members", json={"username": "engcheck", "password": "Eng@12345", "role": "engineer"})
check("دعوة مهندس موقع (دور جديد)", r.status_code == 200, r.text[:120])
c_eng = TestClient(app)
c_eng.post("/api/login", json={"username": "engcheck", "password": "Eng@12345"})
check("المهندس محجوب عن العروض والأسعار (403)",
      c_eng.get("/api/proposals").status_code == 403 and c_eng.get("/api/prices").status_code == 403)
_sub_id = next(o["id"] for o in _opts if "فحص الدرع" in o["label"])
r = c_eng.post("/api/execution/reports", json={
    "project_id": _xp, "report_date": "2026-09-01",
    "lines": [{"item": "صب خرسانة", "unit": "م3", "qty": 40,
               "executor_type": "subcontractor", "subcontractor_id": _sub_id},
              {"item": "أعمال حفر", "unit": "م3", "qty": 65, "executor_type": "company"}]})
check("المهندس يسجّل تقرير إنتاجية بمنفّذين مختلفين", r.status_code == 200, r.text[:120])
_xr = r.json()["id"]
_rep = c_eng.get(f"/api/execution/reports/{_xr}").json()
check("سطر التقرير يحمل اسم المنفّذ", _rep["lines"][0]["executor_name"] == "مؤسسة فحص الدرع"
      and _rep["lines"][1]["executor_name"] == "عمالة الشركة")
r = c_eng.put(f"/api/execution/reports/{_xr}", json={
    "project_id": _xp, "lines": [dict(_rep["lines"][0], qty=45), _rep["lines"][1]]})
check("القفل: تعديل المهندس يتحول لطلب موافقة", r.status_code == 200 and r.json().get("pending"))
_req = r.json().get("request_id")
check("الكمية لم تتغير قبل قرار المالك",
      c_eng.get(f"/api/execution/reports/{_xr}").json()["lines"][0]["qty"] == 40)
_n = c.get("/api/notifications").json()
check("إشعار للمالك بطلب التعديل", any("طلب تعديل" in i["title"] for i in _n["items"]))
r = c.post(f"/api/execution/edit-requests/{_req}/decision", json={"action": "approve"})
check("قبول المالك يطبّق التعديل",
      r.status_code == 200
      and c_eng.get(f"/api/execution/reports/{_xr}").json()["lines"][0]["qty"] == 45)
check("إشعار المهندس بالقبول",
      any("قُبل" in i["title"] for i in c_eng.get("/api/notifications").json()["items"]))
_ps = c.get(f"/api/execution/productivity?project_id={_xp}").json()
_by = {s["executor"]: s for s in _ps}
check("كم اشتغل كل منفّذ: التفصيل بالوحدة صحيح",
      {"unit": "م3", "qty": 45.0} in _by.get("مؤسسة فحص الدرع", {}).get("by_unit", [])
      and {"unit": "م3", "qty": 65.0} in _by.get("عمالة الشركة", {}).get("by_unit", []),
      str(_ps)[:200])

# ---------- الخلاصة ----------
passed = sum(1 for _, ok, _ in RESULTS if ok)
failed = [(n, d) for n, ok, d in RESULTS if not ok]
print(f"\n===== النتيجة: {passed}/{len(RESULTS)} =====")
for n, d in failed:
    print(f"FAILED: {n} — {d}")
sys.exit(1 if failed else 0)
