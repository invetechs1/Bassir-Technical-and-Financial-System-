"""استخراج النصوص من ملفات المشروع المرفوعة (PDF / DOCX / XLSX / TXT).

الملفات الممسوحة ضوئياً (سكان): يُستخدم OCR تلقائياً إن كانت أدواته مثبتة
على الخادم: apt install tesseract-ocr tesseract-ocr-ara poppler-utils
"""
import re
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path

MAX_CHARS_PER_FILE = 60_000
OCR_MARKER = "[استخراج ضوئي OCR]"


def ocr_available() -> bool:
    return bool(shutil.which("pdftoppm") and shutil.which("tesseract"))


def extract_text(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            text = _extract_pdf(content)
        elif ext == ".docx":
            text = _extract_docx(content)
        elif ext in (".xlsx", ".xlsm"):
            text = _extract_xlsx(content)
        elif ext in (".txt", ".md", ".csv"):
            text = content.decode("utf-8", errors="replace")
        else:
            return f"[تنسيق غير مدعوم: {ext}]"
    except Exception as exc:  # ملف تالف أو محمي — نُبلغ بدل أن نفشل
        return f"[تعذر استخراج النص من {filename}: {exc}]"
    text = text.strip()
    if len(text) > MAX_CHARS_PER_FILE:
        text = text[:MAX_CHARS_PER_FILE] + "\n...[تم اقتطاع بقية الملف]"
    return text


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    # نص شبه معدوم = ملف مصور (سكان) — جرّب الاستخراج الضوئي
    if len(re.sub(r"\s", "", text)) < 150:
        ocr_text = _ocr_pdf(content)
        if len(re.sub(r"\s", "", ocr_text)) > len(re.sub(r"\s", "", text)):
            return f"{OCR_MARKER}\n{ocr_text}"
    return text


def _ocr_pdf(content: bytes, max_pages: int = 15) -> str:
    """استخراج ضوئي عبر tesseract (عربي + إنجليزي) — يعمل فقط إن كانت الأدوات مثبتة."""
    if not ocr_available():
        return ""
    try:
        with tempfile.TemporaryDirectory() as td:
            pdf = Path(td) / "in.pdf"
            pdf.write_bytes(content)
            subprocess.run(
                ["pdftoppm", "-r", "200", "-png", "-l", str(max_pages), str(pdf), str(Path(td) / "pg")],
                check=True, capture_output=True, timeout=300,
            )
            langs = "ara+eng" if _has_arabic_ocr() else "eng"
            parts = []
            for img in sorted(Path(td).glob("pg*.png")):
                r = subprocess.run(
                    ["tesseract", str(img), "stdout", "-l", langs, "--psm", "6"],
                    capture_output=True, timeout=120,
                )
                parts.append(r.stdout.decode("utf-8", errors="replace"))
            return "\n".join(parts)
    except Exception:
        return ""


def _has_arabic_ocr() -> bool:
    try:
        r = subprocess.run(["tesseract", "--list-langs"], capture_output=True, timeout=15)
        return b"ara" in r.stdout
    except Exception:
        return False


def _extract_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"## ورقة: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
